import json
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color_hex="CBD5E1"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def build_docx():
    doc = Document()

    # Set page margins (1 inch / 72pt)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Styling helpers
    PRIMARY_COLOR = RGBColor(30, 58, 138)   # #1e3a8a Navy Blue
    SECONDARY_COLOR = RGBColor(37, 99, 235) # #2563eb Blue
    DARK_TEXT = RGBColor(15, 23, 42)       # #0f172a
    MUTED_TEXT = RGBColor(100, 116, 139)   # #64748b

    # Base Normal Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = DARK_TEXT

    # Header / Title Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("CATÁLOGO OFICIAL DE ASIGNATURAS Y UACs")
    run_title.bold = True
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run("BACHILLERATOS GENERALES — MARCO CURRICULAR COMÚN (MCCEMS 2025-2026 / 2026-2027)\nSupervisión Escolar de Educación Media Superior")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = SECONDARY_COLOR

    # Intro box / note
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(14)
    r_intro = p_intro.add_run("📌 Documento Normativo Institucional: Este catálogo concentra de manera organizada y detallada todas las Unidades de Aprendizaje Curricular (UACs) del mapa curricular de Bachillerato General en el estado de Puebla. Está estructurado por componentes y semestres para alimentar al Asistente Virtual IA y servir de referencia oficial para directores, docentes y supervisores.")
    r_intro.font.size = Pt(9.5)
    r_intro.font.italic = True
    r_intro.font.color.rgb = MUTED_TEXT

    # -------------------------------------------------------------
    # RESUMEN EJECUTIVO
    # -------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("📊 Resumen General del Plan Curricular")
    r_h1.font.color.rgb = PRIMARY_COLOR

    t_resumen = doc.add_table(rows=6, cols=3)
    t_resumen.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_resumen)

    headers = ["Sección / Componente Curricular", "Semestres", "Total Asignaturas (UACs)"]
    hdr_cells = t_resumen.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=120, bottom=120)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(10)

    summary_data = [
        ("Sección 1: Currículum Fundamental", "1º a 6º Semestre", "30 UACs"),
        ("Sección 2: Currículum Ampliado (Socioemocionales)", "1º a 6º Semestre", "6 UACs (1 por semestre)"),
        ("Sección 3: Formación Fundamental Extendida Obligatoria (FFEO)", "1º a 6º Semestre", "8 UACs"),
        ("Sección 4: Formación Fundamental Extendida (FFE / Optativas)", "5º y 6º Semestre", "40 UACs (20 por sem)"),
        ("Sección 5: Currículum Laboral (15 Capacitaciones)", "3º a 6º Semestre", "120 UACs (8 por capacitación)")
    ]

    for idx, row_data in enumerate(summary_data):
        row_cells = t_resumen.rows[idx + 1].cells
        bg_hex = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_hex)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if col_idx == 0:
                p.runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECCIÓN 1: CURRÍCULUM FUNDAMENTAL
    # -------------------------------------------------------------
    h_sec1 = doc.add_heading(level=1)
    r_sec1 = h_sec1.add_run("SECCIÓN 1: CURRÍCULUM FUNDAMENTAL (1º a 6º Semestre)")
    r_sec1.font.color.rgb = PRIMARY_COLOR

    p_desc1 = doc.add_paragraph()
    p_desc1.paragraph_format.space_after = Pt(8)
    p_desc1.add_run("El Currículum Fundamental constituye el núcleo formativo esencial del MCCEMS. Consta de 30 UACs distribuidas desde primer hasta sexto semestre en las áreas de Conocimiento y Recursos Sociocognitivos (Lengua y Comunicación, Pensamiento Matemático, Conciencia Histórica, Cultura Digital, Humanidades, Ciencias Naturales y Ciencias Sociales).")

    fundamental_by_sem = {
        1: [
            ("Conciencia Histórica I", "48 hrs", "Recurso Sociocognitivo"),
            ("Cultura Digital I", "48 hrs", "Recurso Sociocognitivo"),
            ("Humanidades I", "64 hrs", "Área de Conocimiento"),
            ("Inglés I", "48 hrs", "Recurso Sociocognitivo"),
            ("La Materia y sus Interacciones", "64 hrs", "Ciencias Naturales"),
            ("Lengua y Comunicación I", "64 hrs", "Recurso Sociocognitivo"),
            ("Pensamiento Matemático I", "64 hrs", "Recurso Sociocognitivo")
        ],
        2: [
            ("Conciencia Histórica II", "48 hrs", "Recurso Sociocognitivo"),
            ("Conservación de la Energía y sus Interacciones con la Materia", "64 hrs", "Ciencias Naturales"),
            ("Cultura Digital II", "48 hrs", "Recurso Sociocognitivo"),
            ("Humanidades II", "64 hrs", "Área de Conocimiento"),
            ("Inglés II", "48 hrs", "Recurso Sociocognitivo"),
            ("Lengua y Comunicación II", "64 hrs", "Recurso Sociocognitivo"),
            ("Pensamiento Matemático II", "64 hrs", "Recurso Sociocognitivo")
        ],
        3: [
            ("Ecosistemas: Interacciones, Energía y Dinámica", "64 hrs", "Ciencias Naturales"),
            ("Humanidades III", "64 hrs", "Área de Conocimiento"),
            ("Inglés III", "48 hrs", "Recurso Sociocognitivo"),
            ("Lengua y Comunicación III", "64 hrs", "Recurso Sociocognitivo"),
            ("Pensamiento Matemático III", "64 hrs", "Recurso Sociocognitivo")
        ],
        4: [
            ("Ciencias Sociales I", "64 hrs", "Área de Conocimiento"),
            ("Conciencia Histórica III", "48 hrs", "Recurso Sociocognitivo"),
            ("Cultura Digital III", "48 hrs", "Recurso Sociocognitivo"),
            ("Formación Socioemocional IV", "32 hrs", "Currículum Ampliado"),
            ("Inglés IV", "48 hrs", "Recurso Sociocognitivo"),
            ("La Superficie Terrestre: Procesos Naturales y Sociales", "64 hrs", "Ciencias Naturales"),
            ("Reacciones Químicas: Conservación de la Materia en la Transformación de la Energía", "64 hrs", "Ciencias Naturales")
        ],
        5: [
            ("Ciencias Sociales II", "64 hrs", "Área de Conocimiento"),
            ("Organismo Vivo: Estructura, Función y Herencia", "64 hrs", "Ciencias Naturales")
        ],
        6: [
            ("Ciencias Sociales III", "64 hrs", "Área de Conocimiento"),
            ("La Biodiversidad y su Conservación", "64 hrs", "Ciencias Naturales")
        ]
    }

    for sem in range(1, 7):
        h_sem = doc.add_heading(level=2)
        r_sem = h_sem.add_run(f"📅 {sem}º Semestre — Currículum Fundamental")
        r_sem.font.color.rgb = SECONDARY_COLOR

        uacs = fundamental_by_sem.get(sem, [])
        t_f = doc.add_table(rows=len(uacs) + 1, cols=4)
        t_f.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t_f)

        f_hdrs = ["#", "Nombre de la Asignatura / UAC", "Horas Totales", "Área / Campo"]
        for i, htext in enumerate(f_hdrs):
            cell = t_f.rows[0].cells[i]
            cell.text = htext
            set_cell_background(cell, "2563EB")
            set_cell_margins(cell, top=100, bottom=100)
            p = cell.paragraphs[0]
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            p.runs[0].font.size = Pt(9.5)

        for u_idx, (u_name, u_hrs, u_area) in enumerate(uacs):
            row_cells = t_f.rows[u_idx + 1].cells
            bg_hex = "F8FAFC" if u_idx % 2 == 1 else "FFFFFF"
            data = [str(u_idx + 1), u_name, u_hrs, u_area]
            for c_i, val in enumerate(data):
                row_cells[c_i].text = val
                set_cell_background(row_cells[c_i], bg_hex)
                set_cell_margins(row_cells[c_i], top=70, bottom=70)
                p = row_cells[c_i].paragraphs[0]
                p.runs[0].font.size = Pt(9)
                if c_i == 1:
                    p.runs[0].font.bold = True

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # SECCIÓN 2: CURRÍCULUM AMPLIADO - FORMACIÓN SOCIOEMOCIONAL
    # -------------------------------------------------------------
    h_sec2 = doc.add_heading(level=1)
    r_sec2 = h_sec2.add_run("SECCIÓN 2: CURRÍCULUM AMPLIADO — FORMACIÓN SOCIOEMOCIONAL (1º a 6º Semestre)")
    r_sec2.font.color.rgb = PRIMARY_COLOR

    p_desc2 = doc.add_paragraph()
    p_desc2.paragraph_format.space_after = Pt(8)
    p_desc2.add_run("El Currículum Ampliado comprende los Recursos Socioemocionales y Ámbitos de Formación Socioemocional. Se imparten 2 horas semanales (32 horas semestrales) desde 1º hasta 6º semestre, abarcando Práctica y Colaboración Ciudadana, Educación para la Salud, Educación Integral en Sexualidad y Género, Actividades Físicas y Deportivas, y Artes.")

    t_soc = doc.add_table(rows=7, cols=4)
    t_soc.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_soc)

    soc_hdrs = ["Semestre", "Nombre Oficial de la UAC", "Horas Semestrales / Totales", "Ámbitos de Formación Socioemocional Incluidos"]
    for i, htext in enumerate(soc_hdrs):
        cell = t_soc.rows[0].cells[i]
        cell.text = htext
        set_cell_background(cell, "2563EB")
        set_cell_margins(cell, top=100, bottom=100)
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9.5)

    soc_data = [
        ("1º Semestre", "Formación Socioemocional I", "2 hrs / sem (32 hrs)", "Práctica y Colaboración Ciudadana, Educación para la Salud, Sexualidad y Género, Deporte y Artes"),
        ("2º Semestre", "Formación Socioemocional II", "2 hrs / sem (32 hrs)", "Práctica y Colaboración Ciudadana, Educación para la Salud, Sexualidad y Género, Deporte y Artes"),
        ("3º Semestre", "Formación Socioemocional III", "2 hrs / sem (32 hrs)", "Práctica y Colaboración Ciudadana, Educación para la Salud, Sexualidad y Género, Deporte y Artes"),
        ("4º Semestre", "Formación Socioemocional IV", "2 hrs / sem (32 hrs)", "Práctica y Colaboración Ciudadana, Educación para la Salud, Sexualidad y Género, Deporte y Artes"),
        ("5º Semestre", "Formación Socioemocional V", "2 hrs / sem (32 hrs)", "Práctica y Colaboración Ciudadana, Educación para la Salud, Sexualidad y Género, Deporte y Artes"),
        ("6º Semestre", "Formación Socioemocional VI", "2 hrs / sem (32 hrs)", "Práctica y Colaboración Ciudadana, Educación para la Salud, Sexualidad y Género, Deporte y Artes")
    ]

    for idx, (s_name, u_name, h_val, amb_val) in enumerate(soc_data):
        row_cells = t_soc.rows[idx + 1].cells
        bg_hex = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
        row_vals = [s_name, u_name, h_val, amb_val]
        for c_i, val in enumerate(row_vals):
            row_cells[c_i].text = val
            set_cell_background(row_cells[c_i], bg_hex)
            set_cell_margins(row_cells[c_i], top=70, bottom=70)
            p = row_cells[c_i].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if c_i == 1:
                p.runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECCIÓN 3: FORMACIÓN FUNDAMENTAL EXTENDIDA OBLIGATORIA (FFEO)
    # -------------------------------------------------------------
    h_sec3 = doc.add_heading(level=1)
    r_sec3 = h_sec3.add_run("SECCIÓN 3: FORMACIÓN FUNDAMENTAL EXTENDIDA OBLIGATORIA (FFEO) (1º a 6º Semestre)")
    r_sec3.font.color.rgb = PRIMARY_COLOR

    p_desc3 = doc.add_paragraph()
    p_desc3.paragraph_format.space_after = Pt(8)
    p_desc3.add_run("La Formación Fundamental Extendida Obligatoria (FFEO) profundiza en la indagación científica, la lectura y redacción avanzada, y el razonamiento matemático. Consta de 8 UACs de carácter obligatorio asignadas de 1º a 6º semestre.")

    t_ffeo = doc.add_table(rows=9, cols=4)
    t_ffeo.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_ffeo)

    ffeo_hdrs = ["Semestre", "Nombre de la Asignatura / UAC", "Clave / Tipo", "Horas Totales"]
    for i, htext in enumerate(ffeo_hdrs):
        cell = t_ffeo.rows[0].cells[i]
        cell.text = htext
        set_cell_background(cell, "2563EB")
        set_cell_margins(cell, top=100, bottom=100)
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9.5)

    ffeo_data = [
        ("1º Semestre", "Laboratorio de Investigación", "ffeo", "64 hrs"),
        ("1º Semestre", "Taller de Lectura y Redacción I", "ffeo", "64 hrs"),
        ("2º Semestre", "Taller de Ciencias I", "ffeo", "64 hrs"),
        ("2º Semestre", "Taller de Lectura y Redacción II", "ffeo", "64 hrs"),
        ("3º Semestre", "Taller de Ciencias II", "ffeo", "64 hrs"),
        ("4º Semestre", "Espacio y Sociedad", "ffeo", "64 hrs"),
        ("5º Semestre", "Taller de Pensamiento Variacional I", "ffeo", "64 hrs"),
        ("6º Semestre", "Temas Selectos de Matemáticas II", "ffeo", "64 hrs")
    ]

    for idx, (s_val, u_val, clv_val, h_val) in enumerate(ffeo_data):
        row_cells = t_ffeo.rows[idx + 1].cells
        bg_hex = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
        row_vals = [s_val, u_val, clv_val, h_val]
        for c_i, val in enumerate(row_vals):
            row_cells[c_i].text = val
            set_cell_background(row_cells[c_i], bg_hex)
            set_cell_margins(row_cells[c_i], top=70, bottom=70)
            p = row_cells[c_i].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if c_i == 1:
                p.runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECCIÓN 4: FORMACIÓN FUNDAMENTAL EXTENDIDA (FFE / OPTATIVAS)
    # -------------------------------------------------------------
    h_sec4 = doc.add_heading(level=1)
    r_sec4 = h_sec4.add_run("SECCIÓN 4: FORMACIÓN FUNDAMENTAL EXTENDIDA (FFE / OPTATIVAS 5º Y 6º SEMESTRE)")
    r_sec4.font.color.rgb = PRIMARY_COLOR

    p_desc4 = doc.add_paragraph()
    p_desc4.paragraph_format.space_after = Pt(8)
    p_desc4.add_run("La Formación Fundamental Extendida (FFE) ofrece 40 asignaturas optativas especializadas (20 en 5º semestre y 20 en 6º semestre) organizadas en áreas de acentuación profesional (Ciencias Naturales, Pensamiento Matemático, Ciencias Sociales, Humanidades y Lenguaje) para preparar al alumno hacia el nivel superior.")

    ffe_5to = [
        ("Análisis de Fenómenos Biológicos", "CNET", "64 hrs"),
        ("Análisis de Fenómenos Físicos I", "CNET", "64 hrs"),
        ("Arte y Cultura I", "Artes/HUM", "64 hrs"),
        ("Comunicación y Sociedad I", "Lenguaje", "64 hrs"),
        ("Derecho y Sociedad I", "Ciencias Sociales", "64 hrs"),
        ("Dibujo Técnico I", "Pensamiento Matemático", "64 hrs"),
        ("Economía I", "Ciencias Sociales", "64 hrs"),
        ("Fundamentos de Administración I", "Ciencias Sociales", "64 hrs"),
        ("Inglés V", "Lenguaje", "64 hrs"),
        ("Lógica y Pensamiento Crítico", "Humanidades", "64 hrs"),
        ("Organización del Flujo de Materia I", "CNET", "64 hrs"),
        ("Pensamiento Filosófico I", "Humanidades", "64 hrs"),
        ("Pensamiento Matemático Finanzas I", "Ciencias Sociales", "64 hrs"),
        ("Probabilidad y Estadística I", "Pensamiento Matemático", "64 hrs"),
        ("Procesos Contables I", "Ciencias Sociales", "64 hrs"),
        ("Psicología I", "Humanidades", "64 hrs"),
        ("Raíces Etimológicas I", "Lenguaje", "64 hrs"),
        ("Salud Integral I", "CNET", "64 hrs"),
        ("Taller Pensamiento Variacional I", "Pensamiento Matemático", "64 hrs"),
        ("Temas Selectos CS I", "Ciencias Sociales", "64 hrs")
    ]

    ffe_6to = [
        ("Análisis de Fenómenos Físicos II", "CNET", "64 hrs"),
        ("Arte y Cultura II", "Artes/HUM", "64 hrs"),
        ("Comunicación y Sociedad II", "Lenguaje", "64 hrs"),
        ("Derecho y Sociedad II", "Ciencias Sociales", "64 hrs"),
        ("Dibujo Técnico II", "Pensamiento Matemático", "64 hrs"),
        ("Economía II", "Ciencias Sociales", "64 hrs"),
        ("Experiencia Estética", "Humanidades", "64 hrs"),
        ("Fundamentos de Administración II", "Ciencias Sociales", "64 hrs"),
        ("Inglés VI", "Lenguaje", "64 hrs"),
        ("Organización del Flujo de Materia II", "CNET", "64 hrs"),
        ("Pensamiento Filosófico II", "Humanidades", "64 hrs"),
        ("Pensamiento Matemático Finanzas II", "Ciencias Sociales", "64 hrs"),
        ("Probabilidad y Estadística II", "Pensamiento Matemático", "64 hrs"),
        ("Procesos Contables II", "Ciencias Sociales", "64 hrs"),
        ("Psicología II", "Humanidades", "64 hrs"),
        ("Raíces Etimológicas II", "Lenguaje", "64 hrs"),
        ("Salud Integral II", "CNET", "64 hrs"),
        ("Taller Pensamiento Variacional II", "Pensamiento Matemático", "64 hrs"),
        ("Temas Selectos CS II", "Ciencias Sociales", "64 hrs"),
        ("Temas Selectos de Biología", "CNET", "64 hrs")
    ]

    for sem_num, ffe_list in [(5, ffe_5to), (6, ffe_6to)]:
        h_fsem = doc.add_heading(level=2)
        r_fsem = h_fsem.add_run(f"📅 {sem_num}º Semestre — Asignaturas FFE Optativas (20 UACs)")
        r_fsem.font.color.rgb = SECONDARY_COLOR

        t_ffe = doc.add_table(rows=len(ffe_list) + 1, cols=4)
        t_ffe.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t_ffe)

        ffe_hdrs = ["#", "Nombre de la Asignatura FFE", "Área de Acentuación / Campo", "Horas Totales"]
        for i, htext in enumerate(ffe_hdrs):
            cell = t_ffe.rows[0].cells[i]
            cell.text = htext
            set_cell_background(cell, "2563EB")
            set_cell_margins(cell, top=100, bottom=100)
            p = cell.paragraphs[0]
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            p.runs[0].font.size = Pt(9.5)

        for u_idx, (u_name, u_area, u_hrs) in enumerate(ffe_list):
            row_cells = t_ffe.rows[u_idx + 1].cells
            bg_hex = "F8FAFC" if u_idx % 2 == 1 else "FFFFFF"
            data = [str(u_idx + 1), u_name, u_area, u_hrs]
            for c_i, val in enumerate(data):
                row_cells[c_i].text = val
                set_cell_background(row_cells[c_i], bg_hex)
                set_cell_margins(row_cells[c_i], top=70, bottom=70)
                p = row_cells[c_i].paragraphs[0]
                p.runs[0].font.size = Pt(9)
                if c_i == 1:
                    p.runs[0].font.bold = True

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # SECCIÓN 5: CURRÍCULUM LABORAL (15 CAPACITACIONES)
    # -------------------------------------------------------------
    h_sec5 = doc.add_heading(level=1)
    r_sec5 = h_sec5.add_run("SECCIÓN 5: CURRÍCULUM LABORAL — CAPACITACIONES (3º a 6º Semestre)")
    r_sec5.font.color.rgb = PRIMARY_COLOR

    p_desc5 = doc.add_paragraph()
    p_desc5.paragraph_format.space_after = Pt(8)
    p_desc5.add_run("El Currículum Laboral abarca las 15 Capacitaciones Oficiales para el Trabajo de Bachillerato General. Cada capacitación se imparte desde 3º hasta 6º semestre (2 submódulos/UACs por semestre, 64 horas cada una). A continuación se presentan las 15 capacitaciones completas desglosadas por semestre:")

    with open(r'C:\NotebookLM\documentos_referencia\Horarios\laboral_grouped.json', 'r', encoding='utf-8') as f:
        laboral_dict = json.load(f)

    for cap_idx, (cap_name, cap_sem_data) in enumerate(laboral_dict.items(), 1):
        h_cap = doc.add_heading(level=2)
        r_cap = h_cap.add_run(f"🛠️ Capacitación {cap_idx}: {cap_name}")
        r_cap.font.color.rgb = SECONDARY_COLOR

        rows_count = 1
        for s_str in ["3", "4", "5", "6"]:
            rows_count += len(cap_sem_data.get(s_str, []))

        t_lab = doc.add_table(rows=rows_count, cols=4)
        t_lab.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t_lab)

        l_hdrs = ["Semestre", "# Submódulo", "Nombre de la UAC / Submódulo Laboral", "Horas Totales"]
        for i, htext in enumerate(l_hdrs):
            cell = t_lab.rows[0].cells[i]
            cell.text = htext
            set_cell_background(cell, "1E3A8A")
            set_cell_margins(cell, top=100, bottom=100)
            p = cell.paragraphs[0]
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            p.runs[0].font.size = Pt(9.5)

        curr_row = 1
        for s_str in ["3", "4", "5", "6"]:
            submod_list = cap_sem_data.get(s_str, [])
            for sub_idx, sub_name in enumerate(submod_list, 1):
                row_cells = t_lab.rows[curr_row].cells
                bg_hex = "F8FAFC" if curr_row % 2 == 1 else "FFFFFF"
                row_vals = [f"{s_str}º Semestre", f"Submódulo {sub_idx}", sub_name, "64 hrs"]
                for c_i, val in enumerate(row_vals):
                    row_cells[c_i].text = val
                    set_cell_background(row_cells[c_i], bg_hex)
                    set_cell_margins(row_cells[c_i], top=70, bottom=70)
                    p = row_cells[c_i].paragraphs[0]
                    p.runs[0].font.size = Pt(9)
                    if c_i == 2:
                        p.runs[0].font.bold = True
                curr_row += 1

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Save outputs
    out_path_1 = r"C:\NotebookLM\documentos_referencia\Horarios\Catalogo_Oficial_Asignaturas_Bachilleratos_Generales_2025-2026.docx"
    out_path_2 = r"C:\Users\samue\.gemini\antigravity-ide\brain\7569d40a-c01f-4ce4-836e-6314c3c5f299\Catalogo_Oficial_Asignaturas_Bachilleratos_Generales_2025-2026.docx"

    doc.save(out_path_1)
    doc.save(out_path_2)

    print(f"Successfully saved docx to {out_path_1} and {out_path_2}")

if __name__ == '__main__':
    build_docx()
